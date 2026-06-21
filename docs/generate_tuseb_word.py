"""
TÜSEB C Grubu RenaCare Proje Başvurusu — Word belgesi üretici
python3 generate_tuseb_word.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── RENK PALETİ ──────────────────────────────────────────────────────────────
NAVY     = RGBColor(0x1e, 0x3a, 0x5f)   # koyu lacivert — başlıklar
TEAL     = RGBColor(0x0d, 0x9a, 0x88)   # teal — vurgu çizgiler
GRAY_DK  = RGBColor(0x2d, 0x3a, 0x4a)   # koyu gri — gövde metin
GRAY_LT  = RGBColor(0xf0, 0xf4, 0xf8)   # açık gri — tablo arka planı
WHITE    = RGBColor(0xff, 0xff, 0xff)
GOLD     = RGBColor(0xd4, 0xa0, 0x17)   # altın — aksan

FONT_NAME = "Calibri"

# ─── YARDIMCI FONKSİYONLAR ────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Hücre arka plan rengini ayarla."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        val = kwargs.get(edge, {})
        if val:
            tag = OxmlElement(f"w:{ edge}")
            for k, v in val.items():
                tag.set(qn(f"w:{k}"), v)
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=11,
            color=None, font=FONT_NAME, underline=False):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def add_para(doc, text="", style="Normal", align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6, line_spacing=None):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    if line_spacing:
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing       = Pt(line_spacing)
    if text:
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY_DK
    return p

def section_heading(doc, number, title, level=1):
    """Numaralı bölüm başlığı — lacivert çizgili."""
    p   = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(16 if level == 1 else 10)
    fmt.space_after  = Pt(4)
    fmt.keep_with_next = True

    if level == 1:
        add_run(p, f"{number}  ", bold=True, size=14, color=TEAL)
        add_run(p, title,         bold=True, size=14, color=NAVY)
        # çizgi ekle
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "4")
        bot.set(qn("w:color"), "0D9A88")
        pBdr.append(bot)
        pPr.append(pBdr)
    else:
        add_run(p, f"{number}  ", bold=True, size=12, color=TEAL)
        add_run(p, title,         bold=True, size=12, color=NAVY)
    return p

def sub_heading(doc, title, level=3):
    p   = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(8)
    fmt.space_after  = Pt(2)
    add_run(p, title, bold=True, size=11, color=NAVY)
    return p

def bullet(doc, text, level=0):
    indent = Cm(0.6 * (level + 1))
    p   = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before     = Pt(1)
    fmt.space_after      = Pt(1)
    fmt.left_indent      = indent
    fmt.first_line_indent= Cm(-0.4)
    add_run(p, "• ", bold=True, size=11, color=TEAL)
    add_run(p, text, size=11, color=GRAY_DK)
    return p

def screenshot_box(doc, number, desc):
    """Ekran görüntüsü yer tutucu kutu."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "EEF4FB")
    set_cell_border(cell,
        top=   {"val":"single","sz":"12","color":"0D9A88"},
        bottom={"val":"single","sz":"12","color":"0D9A88"},
        left=  {"val":"single","sz":"12","color":"0D9A88"},
        right= {"val":"single","sz":"12","color":"0D9A88"},
    )
    cell.width = Cm(14)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, f"[ EKRAN GÖRÜNTÜSÜ {number} ]", bold=True, size=11, color=TEAL)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(10)
    add_run(p2, desc, italic=True, size=10, color=GRAY_DK)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def styled_table(doc, headers, rows, col_widths=None):
    """Şık başlıklı tablo."""
    n_cols = len(headers)
    tbl    = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = "Table Grid"

    # Başlık satırı
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "1e3a5f")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        add_run(p, h, bold=True, size=10, color=WHITE)
        if col_widths:
            cell.width = Cm(col_widths[i])

    # Veri satırları
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg  = "FFFFFF" if ri % 2 == 0 else "F0F4F8"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            if ci == 0:
                add_run(p, val, bold=True, size=10, color=NAVY)
            else:
                add_run(p, val, size=10, color=GRAY_DK)
            if col_widths:
                cell.width = Cm(col_widths[ci])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def info_box(doc, label, text):
    """Vurgulu bilgi kutusu."""
    tbl  = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c0, c1 = tbl.rows[0].cells
    c0.width = Cm(0.4)
    set_cell_bg(c0, "0D9A88")
    set_cell_bg(c1, "EEF4FB")
    p = c1.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    if label:
        add_run(p, label + " ", bold=True, size=10, color=NAVY)
    add_run(p, text, size=10, color=GRAY_DK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def page_break(doc):
    doc.add_page_break()

# ─── BELGE OLUŞTUR ─────────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Sayfa kenar boşlukları
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.5)

    # ── KAPAK SAYFASI ─────────────────────────────────────────────────────────

    # Üst bant
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "1e3a5f")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, "T.C. TÜRKİYE SAĞLIK ENSTİTÜLERİ BAŞKANLIĞI", bold=True, size=13, color=WHITE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(16)
    add_run(p2, "KLİNİK ARAŞTIRMA PROJESİ BAŞVURUSU — C GRUBU", bold=False, size=11, color=RGBColor(0xb0,0xd0,0xee))

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # Logo yerine sembolik kutu
    tbl2 = doc.add_table(rows=1, cols=1)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    cc = tbl2.cell(0, 0)
    set_cell_bg(cc, "EEF4FB")
    cc.width = Cm(5)
    pp = cc.paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_before = Pt(20)
    pp.paragraph_format.space_after  = Pt(20)
    add_run(pp, "♥  RenaCare", bold=True, size=22, color=NAVY)

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # Proje başlığı
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt.paragraph_format.space_before = Pt(4)
    pt.paragraph_format.space_after  = Pt(12)
    add_run(pt,
        "Böbrek Nakli Alıcılarında Yapay Zekâ Destekli\n"
        "Çok Disiplinli Dijital Takip Platformunun (RenaCare)\n"
        "Klinik Etkinliği ve Sağlık Hizmetleri Kullanımına Etkisi:\n"
        "Çok Merkezli, Randomize Kontrollü Çalışma",
        bold=True, size=15, color=NAVY)

    # Alt başlık
    ps = doc.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps.paragraph_format.space_after = Pt(4)
    add_run(ps, "Kısa Ad: RenaCare-RCT", italic=True, size=11, color=GRAY_DK)

    doc.add_paragraph().paragraph_format.space_after = Pt(30)

    # Meta bilgi tablosu
    meta = [
        ("Program",        "TÜSEB Klinik Araştırma Projeleri — C Grubu"),
        ("Çalışma Tipi",   "Çok merkezli, paralel kollu, randomize kontrollü çalışma"),
        ("Önerilen Süre",  "24 ay"),
        ("Başvuru Tarihi", "Haziran 2026"),
        ("Yürütücü Kurum", "[Nakil merkezi olan üniversite/EAH — doldurulacak]"),
        ("Proje Yürütücüsü","[Ad Soyad, Unvan, Kurum]"),
        ("Web / Demo",     "kidneytransplant.vercel.app"),
    ]
    styled_table(doc,
        headers=["Bilgi Alanı", "İçerik"],
        rows=meta,
        col_widths=[5, 11])

    page_break(doc)

    # ── İÇİNDEKİLER ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "İÇİNDEKİLER", bold=True, size=14, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    toc_items = [
        ("1", "Proje Özeti (Türkçe + İngilizce)", "3"),
        ("2", "Bilimsel Gerekçe ve Literatür", "4"),
        ("3", "Platformun Mevcut Durumu ve Özgünlüğü", "6"),
        ("4", "Amaç ve Araştırma Soruları", "10"),
        ("5", "Materyal ve Yöntem", "11"),
        ("6", "Çalışma Takvimi", "14"),
        ("7", "Bütçe Planı", "15"),
        ("8", "Beklenen Çıktılar ve Etki", "16"),
        ("9", "Etik ve Yasal Çerçeve", "17"),
        ("10","Araştırmacı Bilgileri", "17"),
        ("—", "Kaynaklar", "18"),
        ("—", "Ek: Ekran Görüntüleri Rehberi", "19"),
    ]
    for num, title, pg in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        add_run(p, f"{num:>2}.  ", bold=True, size=11, color=TEAL)
        add_run(p, title, size=11, color=GRAY_DK)
        tab_stop = p.paragraph_format
        add_run(p, f"  {'.'*60}  {pg}", size=10, color=RGBColor(0xa0,0xa0,0xa0))

    page_break(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # BÖLÜM 1 — PROJE ÖZETİ
    # ═════════════════════════════════════════════════════════════════════════
    section_heading(doc, "1", "PROJE ÖZETİ")

    sub_heading(doc, "1.1  Türkçe Özet")
    add_para(doc, space_after=4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_run(p,
        "Böbrek nakli, son dönem böbrek yetmezliğinin en etkili tedavisidir; ancak nakil sonrası uzun "
        "dönemde greft sağkalımı ve hasta yaşam kalitesi büyük ölçüde immünosupresif ilaç uyumuna, düzenli "
        "laboratuvar izlemine, beslenme yönetimine ve erken komplikasyon tespitine bağlıdır. Mevcut klinik "
        "izlem sistemleri hasta eğitimi ve öz-yönetim desteği açısından önemli boşluklar barındırmaktadır.\n\n"
        "Bu proje, böbrek nakli alıcılarında ilaç uyumu, greft fonksiyon takibi, beslenme yönetimi ve "
        "hasta-klinisyen iletişimini tek platformda bütünleştiren yapay zekâ destekli dijital sistem "
        "RenaCare'in klinik etkinliğini çok merkezli randomize kontrollü çalışma (RCT) ile sınamayı "
        "amaçlamaktadır. En az 3 nakil merkezinden toplam 120 yetişkin böbrek nakli alıcısı (deney n=60, "
        "kontrol n=60) dahil edilecektir. Birincil sonlanım noktası 6. ayda immünosupresif ilaç uyumsuzluğu "
        "sıklığı (BAASIS ölçeği); ikincil sonlanım noktaları serum kreatinin değişkenliği, tacrolimus düzey "
        "kararlılığı, acil başvuru ve hastane yatış sıklığı, hastayla ilişkili sonuç göstergeleri ve sağlık "
        "hizmetleri maliyetidir. Çalışmadan elde edilecek veriler, dijital sağlık uygulamalarının nakil "
        "kliniklerinde yaygınlaştırılmasına yönelik kanıt tabanı oluşturacak ve klinik karar destek "
        "algoritmaları için gerçek dünya verisi sağlayacaktır.",
        size=11, color=GRAY_DK)

    # Anahtar kelimeler kutusu
    info_box(doc, "Anahtar Kelimeler:",
             "böbrek nakli · ilaç uyumu · dijital sağlık · yapay zekâ · randomize kontrollü çalışma · greft sağkalımı · öz-yönetim")

    sub_heading(doc, "1.2  English Abstract")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_run(p,
        "Kidney transplantation is the most effective treatment for end-stage renal disease; however, "
        "long-term graft survival and patient quality of life depend critically on immunosuppressive medication "
        "adherence, laboratory monitoring, nutritional management, and early complication detection. Existing "
        "clinical monitoring systems have significant gaps in patient education and self-management support.\n\n"
        "This project aims to evaluate the clinical efficacy of RenaCare, an AI-assisted digital system "
        "integrating medication adherence support, graft function monitoring, nutritional management, and "
        "patient-clinician communication for kidney transplant recipients, through a multicenter randomized "
        "controlled trial. A minimum of 120 adult kidney transplant recipients from at least 3 transplant "
        "centers (intervention n=60, control n=60) will be enrolled. The primary outcome is immunosuppressive "
        "medication non-adherence frequency at month 6 (BAASIS scale); secondary outcomes include serum "
        "creatinine variability, tacrolimus trough stability, emergency visits, hospitalizations, "
        "patient-reported outcomes, and healthcare costs.",
        size=11, color=GRAY_DK)

    info_box(doc, "Keywords:",
             "kidney transplantation · medication adherence · digital health · artificial intelligence · "
             "randomized controlled trial · graft survival · self-management")

    page_break(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # BÖLÜM 2 — BİLİMSEL GEREKÇE
    # ═════════════════════════════════════════════════════════════════════════
    section_heading(doc, "2", "BİLİMSEL GEREKÇE VE LİTERATÜR")

    sub_heading(doc, "2.1  Epidemiyolojik Bağlam")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_run(p,
        "Türkiye'de Sağlık Bakanlığı verilerine göre 2024 itibarıyla 25.043 kişi böbrek nakli beklemekte "
        "olup bu sayı, nakil bekleyen hastaların yaklaşık %75'ini oluşturmaktadır. Yılda gerçekleştirilen "
        "böbrek nakli sayısı artmakta; buna karşın nakil sonrası uzun dönemde greft sağkalımı oranları "
        "istenen düzeye ulaşamamaktadır. 5 yıllık canlı verici greft sağkalımı yaklaşık %85–90, kadavradan "
        "nakil için %75–80 iken 10 yıllık veriler belirgin düşüşler sergilemektedir (USRDS 2023). Kronik "
        "greft disfonksiyonunun önde gelen nedenleri arasında ilaç uyumsuzluğu ve geç komplikasyonların "
        "gecikmeli tespiti yer almaktadır.",
        size=11, color=GRAY_DK)

    sub_heading(doc, "2.2  İlaç Uyumsuzluğunun Klinik Önemi")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p,
        "Böbrek nakli alıcılarında immünosupresif ilaç uyumsuzluğu oranları %20–68 arasında değişmekte; "
        "uyumsuzluğun greft kaybı riskini yaklaşık 7 kat artırdığı sistematik derlemelerde "
        "gösterilmektedir (Dew ve ark. 2007). Temel belirleyiciler şöyle sıralanabilir:",
        size=11, color=GRAY_DK)
    for b in ["Bilgi eksikliği ve yetersiz hasta eğitimi",
              "Kompleks çok ilaçlı rejimler",
              "Yan etki korkusu ve tacrolimus tolerabilite sorunları",
              "Günlük rutin entegrasyon güçlüğü",
              "Nakil sonrası izlem desteğinin süreksizliği"]:
        bullet(doc, b)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    sub_heading(doc, "2.3  Dijital Sağlık Müdahaleleri: Kanıt Özeti")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p,
        "Sistematik derleme ve meta-analizler, IT-tabanlı ve mobil sağlık müdahalelerinin nakil "
        "hastalarında ilaç uyumu üzerine olumlu etki potansiyeli taşıdığını göstermektedir. Seçilmiş "
        "kilit çalışmalar:",
        size=11, color=GRAY_DK)

    styled_table(doc,
        headers=["Çalışma", "Tasarım", "Sonuç"],
        rows=[
            ("Dobson ve ark. 2021 (Cochrane)", "Sistematik derleme", "mHealth→iyileşme olası; kanıt yetersiz"),
            ("McGillicuddy 2015 (JMIR)",        "RCT, n=200",          "İlaç uyumu + KB kontrolünde anlamlı iyileşme"),
            ("Kriszbacher ve ark. 2020",         "Sistematik derleme", "Elektronik hatırlatıcı tacrolimus kararlılığını artırır"),
            ("Süzer-Özkan ve ark. 2023",         "Türkiye — tek merkez","Dijital hatırlatıcı bilgi düzeyini artırıyor"),
        ],
        col_widths=[6, 4, 6.5])

    info_box(doc, "Kanıt Boşluğu:",
             "Türk nakil hasta popülasyonuna özgü, çok merkezli, randomize kontrollü ve uzun dönemli "
             "klinik sonlanım noktaları kullanan dijital sağlık müdahale çalışması henüz yayımlanmamıştır. "
             "RenaCare-RCT bu boşluğu doğrudan hedeflemektedir.")

    sub_heading(doc, "2.4  YZ Kullanımında Güvenlik Çerçevesi")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_run(p,
        "RenaCare'in yapay zekâ modülü tanı ya da tedavi önerisi sunmaz; yalnızca doktor onaylı, "
        "böbrek nakli sonrası bakıma özgü eğitim içeriği sunar ve semptom triyajında \"doktorunuzu "
        "arayın\" yönlendirmesiyle kritik uyarılar verir. Bu model, TITCK Tıbbi Cihaz Yazılımları "
        "Kılavuzu ve güncel FDA/CE rehberleriyle uyumlu şeffaf, sınırlı amaçlı YZ kullanımını temsil "
        "etmektedir.",
        size=11, color=GRAY_DK)

    page_break(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # BÖLÜM 3 — PLATFORM
    # ═════════════════════════════════════════════════════════════════════════
    section_heading(doc, "3", "PLATFORMUN MEVCUT DURUMU VE ÖZGÜNLÜĞÜ")

    sub_heading(doc, "3.1  Sistem Mimarisi")
    styled_table(doc,
        headers=["Bileşen", "Teknoloji"],
        rows=[
            ("Frontend",    "Next.js 16 (App Router), TypeScript, Tailwind CSS"),
            ("Veri Katmanı","localStorage (prototip) → KVKK uyumlu şifreli bulut (araştırma fazı)"),
            ("YZ Motoru",   "OpenAI GPT-4 + anahtar kelime tabanlı Türkçe yedek motor"),
            ("Kimlik Doğ.", "localStorage tabanlı JWT → OAuth 2.0 (araştırma fazı)"),
            ("Dağıtım",     "Vercel — kidneytransplant.vercel.app"),
            ("Güvenlik",    "Hız sınırlama (20 istek/dk/IP), HTTPS, KVKK tasarım uyumu"),
        ],
        col_widths=[4, 12.5])

    sub_heading(doc, "3.2  Rol Bazlı Modüller")

    # --- Hasta Paneli
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, "A) Hasta Paneli", bold=True, size=11, color=TEAL)
    for b in [
        "Günlük ilaç programı ve doz onay günlüğü (kalıcı, gün serisi takibi)",
        "Haftalık/aylık ilaç uyum grafiği",
        "Laboratuvar trend izlemi: 14 metrik (kreatinin, GFR, tacrolimus, hemoglobin, CRP, potasyum…)",
        "Beslenme rehberi: kısıtlı besin uyarıları, günlük hedefler (potasyum/fosfor/sıvı)",
        "Randevu planlama ve hatırlatıcılar",
        "Doktor/diyetisyen mesajlaşma",
        "YZ destekli sağlık asistanı (klinik onaylı Türkçe içerik)",
    ]:
        bullet(doc, b)

    screenshot_box(doc, 1,
        "Hasta paneli ana ekranı — anlık lab değerleri (kreatinin, tacrolimus, GFR), "
        "bugünkü ilaç hatırlatıcıları ve son mesajlar. Mobil uyumlu navigasyon.")

    # --- Doktor Paneli
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, "B) Doktor Paneli", bold=True, size=11, color=TEAL)
    for b in [
        "Hasta listesi: durum renk kodlaması (Stabil / Uyarı / Kritik), son kreatinin ve tacrolimus özeti",
        "Hasta profil sayfası: 6 sekme — Genel Bakış, İlaçlar, Laboratuvar, Diyet, Mesajlar, Klinik Notlar",
        "Lab girişi: PDF, Excel, e-Nabız metin yapıştırma ile tarihsel veri içe aktarımı",
        "14 metrik, tarihsel zaman serisinde karşılaştırmalı trend grafikleri",
        "İlk reçete: hasta ekleme sırasında ilaç ve diyet planı tanımlanabilir",
    ]:
        bullet(doc, b)

    screenshot_box(doc, 2,
        "Doktor paneli — hasta listesi: durum rozeti (Stabil/Uyarı/Kritik), "
        "kreatinin + tacrolimus özeti, okunmamış mesaj sayacı. 'Detay Sayfası' butonu tam profil sayfasına yönlendirir.")

    screenshot_box(doc, 3,
        "Hasta profil sayfası — Laboratuvar sekmesi: interaktif metrik seçici, "
        "normal aralık bantlı SVG trend grafiği ve tarihsel lab değerleri tablosu.")

    # --- Analitik Paneli
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, "C) Popülasyon Analitik Paneli (Araştırma Modülü)", bold=True, size=11, color=TEAL)
    for b in [
        "1.000 hastalık geriye dönük anonim kohort istatistikleri",
        "12 aylık trend grafikleri: kreatinin, GFR, tacrolimus, hemoglobin, CRP, potasyum",
        "Hasta durum dağılımı: Stabil %67,8 / Uyarı %23,6 / Kritik %8,6",
        "Süreç göstergeleri: ilaç uyumu %88,7 · lab takip %93,4 · diyet planı kapsamı %81,2",
        "12 aylık değişim ve normal aralık karşılaştırma tablosu (14 metrik)",
        "Nakil yılı, yaş grubu, doktora göre filtreleme",
        "TÜSEB/TÜBİTAK araştırma bağlamı notu ve KVKK uyum bildirimi",
    ]:
        bullet(doc, b)

    screenshot_box(doc, 4,
        "Popülasyon İstatistikleri — üst panel: TÜBİTAK 2577/TÜSEB araştırma notu, "
        "filtre çubuğu, 8 özet kart (n=1.000, Ort. Kreatinin, Ort. GFR, "
        "Tacrolimus Hedefte %, İlaç Uyumu, Lab Takip, Diyet Kapsamı, Stabil %).")

    screenshot_box(doc, 5,
        "Popülasyon İstatistikleri — alt panel: 4 LineChart (Kreatinin / GFR / Tacrolimus / "
        "Diğer Belirteçler) ve 12 aylık metrik özeti tablosu (değişim yönü oklı).")

    # --- Diyetisyen + Koordinatör
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, "D) Diyetisyen Paneli", bold=True, size=11, color=TEAL)
    for b in [
        "Hasta için özelleştirilmiş diyet planı oluşturma ve düzenleme",
        "Kalori / protein / potasyum / fosfor / sıvı hedefleri",
        "Kısıtlı besinler listesi (greyfurt/pomelo dahil ilaç-gıda etkileşimi uyarıları)",
        "Öğün bazlı planlayıcı (kahvaltı, öğle, akşam, ara öğün)",
    ]:
        bullet(doc, b)

    screenshot_box(doc, 6,
        "Diyetisyen paneli — diyet planı düzenleme formu: kalori hedefi, "
        "potasyum/fosfor/sıvı limitleri, kısıtlı besin etiketleri ve öğün bazlı planlayıcı.")

    # --- YZ
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, "E) YZ Destekli Sağlık Asistanı", bold=True, size=11, color=TEAL)
    for b in [
        "GPT-4 tabanlı; API anahtarı olmaksızın anahtar kelime tabanlı Türkçe yedek motor çalışır",
        "Kapsam: tacrolimus, kreatinin, beslenme kısıtlamaları, rejeksiyon belirtileri, acil yönlendirme",
        "Dakikada 20 istek hız sınırı (IP başına); \"tanı koymaz\" güvenlik uyarısı",
    ]:
        bullet(doc, b)

    screenshot_box(doc, 7,
        "YZ Sağlık Asistanı — tacrolimus sorusu ve klinik onaylı Türkçe yanıt, "
        "\"doktor bilgilendirme\" uyarısıyla birlikte. Sol panelde konu kategorileri.")

    screenshot_box(doc, 8,
        "Giriş ekranı: demo hesap hızlı seçim paneli (Hasta / Doktor / Diyetisyen / Koordinatör). "
        "Sağ panel: platform güven noktaları.")

    screenshot_box(doc, 9,
        "İlaç takip sayfası: haftalık uyum grafiği (çubuklarla günlük alınan/toplam), "
        "bugünkü program (alındı/bekliyor rozeti), uyum oranı ve gün serisi kartları.")

    screenshot_box(doc, 10,
        "Doktor mesajlaşma — sol: hasta mesajları listesi (aciliyet + okunmamış rozeti); "
        "sağ: yanıt kutusu ve eş zamanlı epikriz/ilaç özeti.")

    sub_heading(doc, "3.3  Özgünlük Karşılaştırması")
    styled_table(doc,
        headers=["Boyut", "RenaCare", "Mevcut Sistemler"],
        rows=[
            ("Çok rol desteği",             "Hasta + Doktor + Diyetisyen + Koordinatör", "Genellikle tek rol"),
            ("Gerçek zamanlı lab trendleri","14 metrik, SVG grafik, normal bant",         "Çoğunda yok"),
            ("YZ asistanı",                 "Güvenlik sınırlı, hekim onaylı içerik",      "Nadir veya kısıtlı"),
            ("Popülasyon analitiği",         "1.000 hastalık kohort istatistikleri",       "Klinik sistemlerde nadir"),
            ("Lab içe aktarım",             "PDF / Excel / e-Nabız metin yapıştırma",     "Çoğunda yok"),
            ("Türkçe klinik içerik",        "Tam yerelleştirilmiş",                       "Sınırlı"),
            ("KVKK uyumu",                  "Tasarım gereği uyumlu",                      "Değişken"),
        ],
        col_widths=[5, 6, 5.5])

    page_break(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # BÖLÜM 4 — AMAÇ
    # ═════════════════════════════════════════════════════════════════════════
    section_heading(doc, "4", "AMAÇ VE ARAŞTIRMA SORULARI")

    sub_heading(doc, "4.1  Birincil Amaç")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_run(p,
        "RenaCare platformunun 6 aylık kullanımının, standart bakımla karşılaştırıldığında, "
        "böbrek nakli alıcılarında immünosupresif ilaç uyumsuzluğu sıklığını (BAASIS ölçeği) "
        "istatistiksel olarak anlamlı biçimde azaltıp azaltmadığını belirlemek.",
        size=11, color=GRAY_DK)

    sub_heading(doc, "4.2  İkincil Amaçlar")
    for i, b in enumerate([
        "Serum kreatinin değişkenliği (SD/CV) ve tacrolimus trough kararlılığını iki grup arasında karşılaştırmak.",
        "Acil servis başvurusu ve planlanmamış hastane yatışı sıklığını iki grup arasında değerlendirmek.",
        "Hastayla ilişkili sonuç göstergelerini (KDQOL-36, Morisky uyum ölçeği) karşılaştırmak.",
        "Platform kullanılabilirliğini ve kabul edilebilirliğini değerlendirmek (SUS, TAM ölçeği).",
        "Sağlık hizmetleri kullanımı maliyetini iki grup arasında karşılaştırmak.",
        "Klinisyen iş akışı memnuniyetini ölçmek.",
    ], 1):
        bullet(doc, f"{i}. {b}")

    sub_heading(doc, "4.3  Hipotezler")
    info_box(doc, "H₀ (Birincil):",
             "6. ayda deney ve kontrol grupları arasında ilaç uyumsuzluğu oranı bakımından istatistiksel olarak anlamlı fark bulunmaz.")
    info_box(doc, "H₁ (Birincil):",
             "6. ayda RenaCare grubunda ilaç uyumsuzluğu oranı kontrol grubuna göre en az %20 düşük olacaktır.")

    page_break(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # BÖLÜM 5 — YÖNTEM
    # ═════════════════════════════════════════════════════════════════════════
    section_heading(doc, "5", "MATERYAL VE YÖNTEM")

    sub_heading(doc, "5.1  Çalışma Tasarımı")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_run(p,
        "Çok merkezli (≥3 merkez), prospektif, paralel kollu (1:1), randomize kontrollü çalışma. "
        "Randomizasyon, nakil tipi (canlı/kadavra) ve nakil sonrası süre (<1 yıl / ≥1 yıl) "
        "tabakalandırmasıyla blok randomizasyon yöntemiyle gerçekleştirilecektir. Müdahalenin niteliği "
        "gereği katılımcı ve klinisyen körlenmesi mümkün değildir; birincil sonlanım noktası "
        "değerlendirmesi kör süreçle (outcome assessor blinding) yürütülecektir.",
        size=11, color=GRAY_DK)

    sub_heading(doc, "5.2  Katılım Kriterleri")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Dahil Edilme Ölçütleri:", bold=True, size=11, color=NAVY)
    for b in ["18 yaş ve üzeri",
              "Böbrek nakli yapılmış olmak (canlı veya kadavra, en az 3 ay önce)",
              "Rutin nefroloji takibinde olmak",
              "İnternet ve akıllı telefon erişimine sahip olmak",
              "Türkçe anlayabilmek", "Yazılı bilgilendirilmiş onam vermek"]:
        bullet(doc, b)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, "Dışlanma Ölçütleri:", bold=True, size=11, color=NAVY)
    for b in ["Aktif RCT'ye katılım veya araştırma müdahalesi",
              "Aktif akut rejeksiyon veya YBÜ gerektiren klinik tablo",
              "Bilişsel yetersizlik (Mini-Cog ≤ 2)",
              "İnternet erişiminin hiç olmaması"]:
        bullet(doc, b)

    sub_heading(doc, "5.3  Örneklem Büyüklüğü")
    styled_table(doc,
        headers=["Varsayım", "Değer"],
        rows=[
            ("Kontrol grubu uyumsuzluk oranı (literatür)", "%45"),
            ("Beklenen mutlak azalma (deney grubu)",        "%20 (→ %25)"),
            ("α (çift yönlü)",                              "0,05"),
            ("Güç (1−β)",                                   "%80"),
            ("Her kolda gerekli n (χ²)",                    "47"),
            ("%15 kayıp/çekilme payı sonrası",              "Her kolda 55"),
            ("Güvenlik payı ile hedef toplam",              "120 (60 + 60)"),
        ],
        col_widths=[10, 6.5])

    sub_heading(doc, "5.4  Müdahale")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Deney Grubu: ", bold=True, size=11, color=TEAL)
    add_run(p,
        "Standart bakım + RenaCare platformu (6 ay, günlük kullanım). İlk hafta klinisyen "
        "refakatinde oryantasyon seansı. Bileşenler: ilaç hatırlatıcısı, lab trend görüntüleme, "
        "beslenme rehberi, randevu takibi, mesajlaşma, YZ eğitim asistanı.",
        size=11, color=GRAY_DK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_run(p, "Kontrol Grubu: ", bold=True, size=11, color=TEAL)
    add_run(p,
        "Standart bakım (rutin poliklinik + nakil sonrası standart hasta eğitimi). "
        "Çalışma sonunda platforma ertelenmiş erişim sağlanacaktır.",
        size=11, color=GRAY_DK)

    sub_heading(doc, "5.5  Ölçüm Araçları")
    styled_table(doc,
        headers=["Sonlanım", "Ölçüm Aracı", "Zaman Noktaları"],
        rows=[
            ("İlaç uyumsuzluğu",    "BAASIS",                   "0 — 3 — 6. ay"),
            ("Yaşam kalitesi",       "KDQOL-36",                 "0 — 6. ay"),
            ("Öz-yönetim",           "KDSM Survey",              "0 — 3 — 6. ay"),
            ("Kullanılabilirlik",    "System Usability Scale",   "1 — 3 — 6. ay (deney)"),
            ("Kabul",               "TAM anketi",               "1 — 6. ay (deney)"),
            ("Lab verisi",          "Rutin klinik kayıt",       "Aylık"),
            ("Sağlık hizm. kull.",  "Acil / yatış (hastane)",   "6 ay boyunca"),
            ("Klinisyen memnun.",   "Likert yapılandırılmış",   "3 — 6. ay"),
        ],
        col_widths=[4.5, 5, 4])

    sub_heading(doc, "5.6  İstatistiksel Analiz")
    for b in [
        "Birincil sonlanım: χ² veya Fisher's exact; odds ratio (%95 GA)",
        "Sürekli değişkenler: bağımsız örneklem t-testi veya Mann-Whitney U",
        "Tekrarlanan ölçüm: MMRM (karışık model tekrarlayan ölçüm ANOVA)",
        "Kayıp veri: çoklu atama (multiple imputation); gerektiğinde LOCF",
        "Niyet-tedavi analizi (ITT) birincil; per-protokol analiz ikincil",
        "Anlamlılık eşiği: p < 0,05 (çift yönlü); çoklu karşılaştırmada Bonferroni",
    ]:
        bullet(doc, b)

    page_break(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # BÖLÜM 6 — TAKVİM
    # ═════════════════════════════════════════════════════════════════════════
    section_heading(doc, "6", "ÇALIŞMA TAKVİMİ (24 AY)")

    styled_table(doc,
        headers=["Faz", "Aylar", "Temel Faaliyetler"],
        rows=[
            ("Hazırlık",          "1–3",   "Etik kurul + KVKK onayları, platform klinik uyarlama, klinisyen eğitimi, pilot test (n=5)"),
            ("Katılım",           "2–6",   "Hasta davet, tabakalı blok randomizasyon, oryantasyon seansları"),
            ("Müdahale & İzlem",  "3–12",  "6 aylık aktif platform kullanımı, aylık lab takibi, 3. ay ara değerlendirme"),
            ("Çekilme Dönemi",    "12",    "Son veri toplama, çalışma kapatma"),
            ("Veri Analizi",      "13–18", "İstatistiksel analiz, nitel veri kodlaması, maliyet analizi"),
            ("Yayım",             "19–24", "Makale yazımı, kongre bildirisi, TÜSEB nihai raporu"),
        ],
        col_widths=[4, 2.5, 10])

    info_box(doc, "Önkoşul Not:",
             "Bu RCT, TÜBİTAK 1002-A kapsamında planlanan ön pilot çalışmanın (n=30, tek kol) "
             "tamamlanmasının ardından güç hesabı ve protokol güncellemesiyle devreye alınacaktır. "
             "İki başvuru birbirini tamamlayıcı niteliktedir: 1002→pilot kanıt, TÜSEB C Grubu→RCT.")

    page_break(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # BÖLÜM 7 — BÜTÇE
    # ═════════════════════════════════════════════════════════════════════════
    section_heading(doc, "7", "BÜTÇE PLANI")

    sub_heading(doc, "7.1  Personel Giderleri")
    styled_table(doc,
        headers=["Pozisyon", "Katkı %", "Ay", "Birim (TL)", "Toplam (TL)"],
        rows=[
            ("Proje Yürütücüsü (Doç/Prof)",        "%30", "24", "15.000",  "108.000"),
            ("Nefroloji Araştırmacısı × 2",         "%20", "24", "12.000",  "115.200"),
            ("Yazılım Müh. / Tıp Bilişimi",         "%60", "24", "20.000",  "288.000"),
            ("Diyetisyen Araştırmacısı × 2",        "%20", "24",  "8.000",   "76.800"),
            ("Klinik Araştırma Koordinatörü",       "%80", "24", "10.000",  "192.000"),
            ("Biyoistatistikçi",                    "%20", "18", "12.000",   "43.200"),
            ("TOPLAM",                              "",    "",   "",        "823.200"),
        ],
        col_widths=[6.5, 2, 1.5, 3, 3.5])

    sub_heading(doc, "7.2  Sarf / Ekipman / Diğer Giderler")
    styled_table(doc,
        headers=["Kalem", "Gerekçe", "Tutar (TL)"],
        rows=[
            ("Bulut altyapısı (KVKK uyumlu)", "24 ay barındırma, SSL, yedekleme",          "72.000"),
            ("OpenAI API erişimi",             "GPT-4 API, 24 ay token kullanımı",          "36.000"),
            ("REDCap / eşdeğer lisans",        "Klinik veri yönetim sistemi",               "18.000"),
            ("Hasta anketi basım/dijital",     "BAASIS, KDQOL-36, SUS uygulama",           "12.000"),
            ("Araştırma bilgisayarı × 2",      "Dizüstü, araştırma ekibi",                  "60.000"),
            ("Yazılım lisansları",             "IDE, test araçları, yıllık",               "12.000"),
            ("Merkez koordinasyon topl. × 6",  "Seyahat, konaklama",                       "30.000"),
            ("Ulusal kongre × 2",              "Kayıt + seyahat",                          "24.000"),
            ("Uluslararası kongre × 1",        "Kayıt + uçak + konaklama",                 "40.000"),
            ("Makale yayın bedeli × 2",        "Open access Q1/Q2 dergi",                  "60.000"),
            ("Çeviri ve redaksiyon",           "İngilizce makale",                         "15.000"),
            ("Proje yönetim giderleri",        "Yazışma, kargo, sarf",                     "18.000"),
            ("ARA TOPLAM (Sarf/Ekip./Diğer)", "",                                          "397.000"),
        ],
        col_widths=[5.5, 7, 4])

    sub_heading(doc, "7.3  Genel Toplam")
    styled_table(doc,
        headers=["Kategori", "Tutar (TL)"],
        rows=[
            ("Personel Giderleri",     "823.200"),
            ("Sarf / Ekipman / Diğer", "397.000"),
            ("GENEL TOPLAM",           "1.220.200"),
        ],
        col_widths=[12, 4.5])

    info_box(doc, "Not:",
             "Yukarıdaki bütçe tahmini niteliğindedir. TÜSEB başvuru sistemindeki güncel birim fiyat "
             "sınırları ve program çerçevesine göre revize edilecektir.")

    page_break(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # BÖLÜM 8 — ÇIKTILAR
    # ═════════════════════════════════════════════════════════════════════════
    section_heading(doc, "8", "BEKLENEN ÇIKTILAR VE ETKİ")

    sub_heading(doc, "8.1  Bilimsel Çıktılar")
    styled_table(doc,
        headers=["Çıktı Türü", "Hedef", "Süre"],
        rows=[
            ("SCI/SCI-E makale (Q1/Q2)",    "≥ 2 tam metin",             "Ay 20–24"),
            ("Ulusal konferans bildirisi",    "≥ 2",                       "Ay 12–18"),
            ("Uluslararası konf. bildirisi",  "≥ 1",                       "Ay 18–24"),
            ("Araştırma protokolü yayımı",   "PROSPERO + protokol makalesi","Ay 3"),
            ("Lisansüstü tez katkısı",       "≥ 1",                       "Ay 24"),
        ],
        col_widths=[6, 5, 5.5])

    sub_heading(doc, "8.2  Klinik ve Toplumsal Etki")
    for b in [
        "İlaç uyumu artışı: %20 iyileşme → doğrudan greft sağkalımı katkısı",
        "Komplikasyon erken tespiti: lab trend uyarıları ile rejeksiyon/toksisite fark edilmesi",
        "Sağlık sistemi tasarrufu: acil başvuru ve planlanmamış yatış azalması",
        "Hasta güçlendirme: öz-yönetim becerisi ve sağlık okuryazarlığı artışı",
        "Çok disiplinli bakım: doktor-diyetisyen-koordinatör-hasta entegrasyonu",
    ]:
        bullet(doc, b)

    sub_heading(doc, "8.3  Politika ve Ölçeklenme")
    for b in [
        "TITCK Tıbbi Cihaz Yazılımları kılavuzu çerçevesinde gerçek dünya klinik kanıt",
        "Nakil polikliniğine entegre edilebilir ölçeklenebilir platform modeli",
        "Çok merkezli uluslararası çalışmalara hazır protokol paketi",
        "Açık kaynak bileşen yayımı: farklı kronik hastalıklar için uyarlama zemini",
    ]:
        bullet(doc, b)

    page_break(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # BÖLÜM 9 — ETİK
    # ═════════════════════════════════════════════════════════════════════════
    section_heading(doc, "9", "ETİK VE YASAL ÇERÇEVE")

    styled_table(doc,
        headers=["Gereksinim", "Durum"],
        rows=[
            ("Girişimsel Olmayan Klinik Araşt. Etik Kurulu",     "Proje onayı sonrasında başvurulacak"),
            ("KVKK Uyumu",                                        "Platform tasarımı gereği uyumlu; Veri İşleme Sözleşmesi hazırlanacak"),
            ("Bilgilendirilmiş Onam",                             "Yazılı, CIOMS kılavuzları uyumlu"),
            ("PROSPERO Tescili",                                   "Çalışma başlamadan tamamlanacak"),
            ("Klinik Araştırmalar Hakkında Yönetmelik",           "Araştırmacı GKİ sertifikaları mevcut / alınacak"),
            ("Veri Güvenliği",                                    "ISO/IEC 27001 uyumlu bulut; kriptografik şifreleme"),
        ],
        col_widths=[7.5, 9])

    # ═════════════════════════════════════════════════════════════════════════
    # BÖLÜM 10 — ARAŞTIRMACI BİLGİLERİ
    # ═════════════════════════════════════════════════════════════════════════
    section_heading(doc, "10", "ARAŞTIRMACI BİLGİLERİ VE GÖREV DAĞILIMI")

    styled_table(doc,
        headers=["Araştırmacı", "Unvan / Kurum", "Görev", "Katkı %"],
        rows=[
            ("[Proje Yürütücüsü]",      "Doç. Dr. — Nefroloji",       "Proje yönetimi, birincil sonlanım analizi",    "%30"),
            ("[Araştırmacı 2]",         "Uzm. Dr. — Nefroloji",       "Hasta katılımı, 2. merkez koordinasyonu",     "%20"),
            ("[Araştırmacı 3]",         "Uzm. Dr. — Nefroloji",       "3. merkez koordinasyonu",                     "%20"),
            ("Fatih Bildirici",          "Yazılım / Tıp Bilişimi",     "Platform geliştirme, klinik uyarlama, veri altyapısı","%60"),
            ("[Araştırmacı 5]",         "Diyetisyen",                 "Beslenme protokolü, içerik onayı",            "%20"),
            ("[Araştırmacı 6]",         "Diyetisyen",                 "2. diyetisyen, hasta eğitimi",                "%20"),
            ("[Araştırmacı 7]",         "Biyoistatistikçi",           "Analiz planı, güç hesabı, raporlama",         "%20"),
            ("[Koordinatör]",            "Klinik Araştırma Koordinatörü","Randomizasyon, veri kalite kontrolü",      "%80"),
        ],
        col_widths=[3.5, 4, 6, 2.5])

    page_break(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # KAYNAKLAR
    # ═════════════════════════════════════════════════════════════════════════
    section_heading(doc, "—", "KAYNAKLAR")

    refs = [
        "TÜSEB. Klinik Araştırma Projeleri Uygulama Kılavuzu. [Güncel erişim bağlantısı eklenecek]",
        "T.C. Sağlık Bakanlığı. Ulusal Organ ve Doku Nakli İstatistikleri. Ankara, 2024.",
        "Dew MA ve ark. Rates and risk factors for nonadherence to the medical regimen after adult solid organ transplantation. Transplantation. 2007;83(7):858–873.",
        "Rosenberger J ve ark. Medication adherence and health-related quality of life in kidney transplant recipients. Transpl Int. 2012.",
        "Denhaerynck K ve ark. Nonadherence to immunosuppressive medication: systematic review. Transpl Int. 2020.",
        "Dobson CT ve ark. IT-based interventions for adherence in kidney transplant recipients. Cochrane Database Syst Rev. 2021.",
        "McGillicuddy JW ve ark. Mobile health decision support for hypertension in transplant patients. Am J Kidney Dis. 2015.",
        "Kriszbacher I ve ark. Electronic reminders for tacrolimus adherence: systematic review. Patient Prefer Adherence. 2020.",
        "Süzer-Özkan FN ve ark. Digital health tools for kidney transplant recipients. Transplant Proc. 2023.",
        "USRDS. 2023 Annual Data Report. NIH/NIDDK, Bethesda MD, 2023.",
        "T.C. Kişisel Verileri Koruma Kurumu. Kişisel Sağlık Verilerine İlişkin Rehber. Ankara: KVKK, 2023.",
        "TITCK. Tıbbi Cihaz Yazılımları Hakkında Kılavuz. Ankara: Türkiye İlaç ve Tıbbi Cihaz Kurumu, 2023.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent      = Cm(0.6)
        p.paragraph_format.first_line_indent= Cm(-0.6)
        p.paragraph_format.space_before     = Pt(2)
        p.paragraph_format.space_after      = Pt(2)
        add_run(p, f"[{i}]  ", bold=True, size=10, color=TEAL)
        add_run(p, ref, size=10, color=GRAY_DK)

    page_break(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # EK — EKRAN GÖRÜNTÜLERİ REHBERİ
    # ═════════════════════════════════════════════════════════════════════════
    section_heading(doc, "EK", "EKRAN GÖRÜNTÜLERİ REHBERİ")

    info_box(doc, "Demo Erişim:",
             "kidneytransplant.vercel.app   —   tüm şifreler: Demo1234")

    styled_table(doc,
        headers=["No", "Ekran", "Demo Hesap", "Bölüm"],
        rows=[
            ("SS-01", "Hasta paneli — lab kartları + ilaç hatırlatıcı",        "hasta@demo.com",        "3.2 A"),
            ("SS-02", "Doktor paneli — hasta listesi (durum renk kodlaması)",    "doktor@demo.com",       "3.2 B"),
            ("SS-03", "Hasta profil — Laboratuvar sekmesi (grafik + tablo)",     "doktor@demo.com",       "3.2 B"),
            ("SS-04", "Popülasyon İstatistikleri — özet kartlar",               "doktor@demo.com",       "3.2 C"),
            ("SS-05", "Popülasyon İstatistikleri — trend grafikleri + tablo",    "doktor@demo.com",       "3.2 C"),
            ("SS-06", "Diyetisyen paneli — diyet planı formu",                  "diyetisyen@demo.com",   "3.2 D"),
            ("SS-07", "YZ Asistan — tacrolimus sorusu + güvenli yanıt",         "hasta@demo.com",        "3.2 E"),
            ("SS-08", "Giriş ekranı — demo hesap seçim paneli",                 "(giriş öncesi)",        "3.3"),
            ("SS-09", "İlaç takip — haftalık uyum grafiği + program",           "hasta@demo.com",        "3.2 A"),
            ("SS-10", "Doktor mesajlaşma — liste + yanıt + epikriz",            "doktor@demo.com",       "3.2 B"),
        ],
        col_widths=[1.8, 7, 4, 2.5])

    # Son not
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p,
        "Bu belge RenaCare TÜSEB C Grubu başvurusu için taslak olarak hazırlanmıştır.\n"
        "Başvuru öncesinde güncel TÜSEB kılavuzu incelenmeli; kurum imzalı taahhüt mektubu,\n"
        "araştırmacı CV'leri ve etik kurul dilekçesi belgeye eklenmelidir.",
        italic=True, size=9, color=RGBColor(0x80,0x90,0xa0))

    return doc

# ─── KAYDET ───────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    out_path = "/Users/fatihb/Downloads/kidneytransplant-main/docs/TUSEB_C_Grubu_RenaCare_Basvurusu.docx"
    doc = build_document()
    doc.save(out_path)
    print(f"✓  Belge oluşturuldu: {out_path}")
